import os
import logging
import pypdf
from docx import Document
import tempfile
from config import Config
from nanonetocr import NanoNetsOCRProcessor
# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)
class DocumentProcessor:
    """Handles document parsing and text extraction with Nanonets preprocessing."""
    
    def __init__(self):
        self.nanonets_processor = NanoNetsOCRProcessor()
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file using Nanonets OCR."""
        try:
            # Use Nanonets OCR for better results
            if self.nanonets_processor.model:
                logger.info(f"🔄 Processing PDF with Nanonets OCR: {file_path}")
                
                # Convert PDF to images
                pages_content = self.nanonets_processor.extract_pages_content(file_path)
                
                if not pages_content:
                    # Fallback to traditional PDF extraction
                    return self._extract_text_from_pdf_traditional(file_path)
                
                # Process each page and combine results
                combined_text = ""
                for i, page_content in enumerate(pages_content):
                    page_text = ""
                    for element in page_content:
                        if element['type'] == 'text':
                            page_text += element['content'] + "\n"
                        elif element['type'] == 'image':    
                            # save to temp file first
                            temp_file = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
                            element['content'].save(temp_file.name)
                            page_text += self.nanonets_processor.convert_to_markdown(temp_file.name) + "\n"
                            os.unlink(temp_file.name)
                            temp_file.close()
                    logger.info(f"📄 Processing page {i+1}/{len(pages_content)}")
                    
                    if page_text:
                        combined_text += f"\n\n--- Page {i+1} ---\n\n{page_text}"
                    
                if combined_text:
                    logger.info(f"✅ Extracted text from PDF using Nanonets: {len(combined_text)} characters")
                    ## save to temp file
                    with open(f"temp_page_{i+1}.md", "w") as file:
                        file.write(combined_text)
                    return combined_text.strip()
            
            # Fallback to traditional extraction
            return self._extract_text_from_pdf_traditional(file_path)
            
        except Exception as e:
            logger.error(f"❌ Error extracting text from PDF: {e}")
            return self._extract_text_from_pdf_traditional(file_path)
    
    def _extract_text_from_pdf_traditional(self, file_path: str) -> str:
        """Traditional PDF text extraction as fallback."""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            
            logger.info(f"✅ Extracted text from PDF (traditional): {len(text)} characters")
            return text.strip()
            
        except Exception as e:
            logger.error(f"❌ Error extracting text from PDF (traditional): {e}")
            return ""
    
    def extract_text_from_image(self, file_path: str) -> str:
        """Extract text from image file using Nanonets OCR."""
        try:
            if self.nanonets_processor.model:
                logger.info(f"🔄 Processing image with Nanonets OCR: {file_path}")
                markdown_text = self.nanonets_processor.convert_to_markdown(file_path)
                
                if markdown_text:
                    logger.info(f"✅ Extracted text from image: {len(markdown_text)} characters")
                    return markdown_text
            
            logger.warning(f"⚠️ Nanonets OCR not available for image: {file_path}")
            return ""
            
        except Exception as e:
            logger.error(f"❌ Error extracting text from image: {e}")
            return ""
    
    def extract_text_from_html(self, file_path: str) -> str:
        """Extract text from HTML file."""
        try:
            return self.nanonets_processor.process_html_file(file_path)
        except Exception as e:
            logger.error(f"❌ Error extracting text from HTML: {e}")
            return ""
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            logger.info(f"✅ Extracted text from DOCX: {len(text)} characters")
            return text.strip()
            
        except Exception as e:
            logger.error(f"❌ Error extracting text from DOCX: {e}")
            return ""
    
    @staticmethod
    def extract_text_from_txt(file_path: str) -> str:
        """Extract text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            logger.info(f"✅ Extracted text from TXT: {len(text)} characters")
            return text.strip()
            
        except Exception as e:
            logger.error(f"❌ Error extracting text from TXT: {e}")
            return ""
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from supported file formats."""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_ext == '.docx':
            return self.extract_text_from_docx(file_path)
        elif file_ext == '.txt':
            return self.extract_text_from_txt(file_path)
        elif file_ext in ['.png', '.jpg', '.jpeg', '.tiff', '.tif', '.bmp', '.webp']:
            return self.extract_text_from_image(file_path)
        elif file_ext in ['.html', '.htm']:
            return self.extract_text_from_html(file_path)
        else:
            logger.error(f"❌ Unsupported file format: {file_ext}")
            return ""

